import os
import json
import tkinter as tk
from tkinter import ttk,colorchooser,filedialog,colorchooser,messagebox,simpledialog
import tkinter.font as tkfont
from docx import Document
from docx.shared import RGBColor
import platformdirs

program_path = os.path.dirname(os.path.abspath(__file__))

        
class EditColorMapping(tk.Toplevel):
    def __init__(self, parent, mapping, colorize_text_callback, update_mapping_callback):
        super().__init__(parent)
        
        # Block parent window
        self.grab_set()
        
        #self.resizable(width=0,height=0)
        self.mapping = mapping.copy()
        self.colorize_text_callback = colorize_text_callback
        self.update_mapping_callback = update_mapping_callback
        
        # Workaround from https://stackoverflow.com/questions/61105126/tag-configure-is-not-working-while-using-theme-ttk-treeview
        self.style = ttk.Style(self)
        aktualTheme = self.style.theme_use()
        if "dummy" not in self.style.theme_names():
            self.style.theme_create("dummy", parent=aktualTheme)
        self.style.theme_use("dummy")

        self.treeview = ttk.Treeview(self, columns=('color'))
        self.treeview.heading('#0', text='Letter')
        self.treeview.heading('color', text='Color')
    
        for letter, color in self.mapping.items():
            self.treeview.insert('', 'end', tag=letter, text=letter, values=(color))
            self.treeview.tag_configure(letter, foreground=color,background="white")
        self.treeview.column('#0', width=200)
        self.treeview.column('color', width=200)

        self.treeview.bind('<Double-1>', self.edit_color)
        self.treeview.bind('<Return>', self.edit_color)
        self.treeview.bind('<Delete>', self.delete_mapping)

        #self.treeview.pack(side='left',expand=True)

        vsb = ttk.Scrollbar(self, orient="vertical", command=self.treeview.yview)
        self.treeview.configure(yscrollcommand=vsb.set)

        button_frame = ttk.Frame(self)
        #button_frame.pack(side='bottom', pady=10)

        add_button = ttk.Button(button_frame, text='Add Mapping', command=self.add_mapping)
        add_button.pack(side='left', padx=5)

        remove_button = ttk.Button(button_frame, text='Remove Mapping', command=self.remove_mapping)
        remove_button.pack(side='left', padx=5)
        
        manual_color_button = ttk.Button(button_frame, text='Edit Color Name', command=self.edit_color_string)
        manual_color_button.pack(side='left', padx=5)
        
        cancel_button = ttk.Button(button_frame, text='Cancel', command=self.cancel)
        cancel_button.pack(side='right', padx=5)

        ok_button = ttk.Button(button_frame, text='OK', command=self.ok)
        ok_button.pack(side='right', padx=5)
        
        self.protocol("WM_DELETE_WINDOW", self.cancel)

        # pack the TreeView, Scrollbar, and button frames using grid
        self.treeview.grid(row=0, column=0, sticky=tk.NSEW)
        vsb.grid(row=0, column=1, sticky=tk.NS)
        button_frame.grid(row=1, column=0, sticky=tk.EW)

    def edit_color(self, event=None):
        item_id = self.treeview.focus()
        letter = self.treeview.item(item_id)['text']
        color = self.treeview.item(item_id)['values'][0]
        new_color = colorchooser.askcolor(color=color, parent=self)
        if new_color[1]:
            self.update_color(item_id, letter, new_color[1])

    def edit_color_string(self, event=None):
        item_id = self.treeview.focus()
        letter = self.treeview.item(item_id)['text']
        color = self.treeview.item(item_id)['values'][0]
        new_color = simpledialog.askstring(
            title=letter, 
            prompt= "Please enter a new color for the letter " + letter + ":", 
            initialvalue = color)
        if new_color is not None:
            try:
                # try to create a 1x1 pixel image with the color string
                image = tk.PhotoImage(width=1, height=1)
                image.put(new_color, to=(0, 0))
                # if the color string is valid, the image is created without error
            except tk.TclError:
                # if the color string is invalid, a TclError is raised
                messagebox.showerror('Error', 'Error: The color code ' + new_color + 
                                     ' does not have the correct format.' +
                                     'Please provide hexcodes as #abcdef.')
                return
            self.update_color(item_id, letter, new_color)

    def update_color(self,item_id,letter,color):          
        self.mapping[letter] = color
        self.treeview.tag_configure(letter, foreground=color, background="white")
        self.treeview.set(item_id, 'color', color)
        self.treeview.item(item_id,tag=letter)
        # Apply updated color map temporarily in main window
        self.colorize_text_callback(self.mapping)
        
    def delete_mapping(self, event=None):
        item_id = self.treeview.focus()
        next_id = self.treeview.next(item_id)
        letter = self.treeview.item(item_id)['text']
        del self.mapping[letter]
        self.treeview.delete(item_id)
        self.treeview.focus(next_id)
        self.colorize_text_callback(self.mapping)

    def add_mapping(self):
        letter = tk.simpledialog.askstring('Add Mapping', 'Enter a letter:', parent=self)
        if letter and letter not in self.mapping:
            self.mapping[letter] = '#ffffff'
            item = self.treeview.insert('', 'end', text=letter, values=(letter, '#ffffff', 'Edit'))
            self.treeview.set(item, 'color', '#ffffff')
            self.colorize_text_callback(self.mapping)
        elif letter in self.mapping:
            messagebox.showerror('Error', 'Error: Letter already in list.')
            
    def remove_mapping(self):
        item_id = self.treeview.focus()
        letter = self.treeview.item(item_id)['text']
        del self.mapping[letter]
        self.treeview.delete(item_id)
                
    def cancel(self):
        self.colorize_text_callback()
        self.destroy()
    
    def ok(self):
        self.update_mapping_callback(self.mapping)
        self.colorize_text_callback()
        self.destroy()
        
class TextEditor:
    def __init__(self, master):
        self.master = master
        master.title("Synaesthesia Writer")
        self.config_file=os.path.join(
            platformdirs.user_config_dir(appname="SynaesthesiaWriter"),
            'synaesthesia_writer_config.json');
        try:
            with open(self.config_file, 'r') as f:
                print('Found config file', self.config_file, ' -> load it')
                self.config = json.load(f)
        except:
            # default color mapping
            default_color_mapping = {
                "a": "red",
                "b": "blue",
                "c": "#2596be",
                "d": "orange",
                "e": "purple",
                "f": "brown",
                "g": "grey",
                "h": "pink",
                "i": "cyan",
                "j": "magenta",
                "k": "navy",
                "l": "olive",
                "m": "navy",
                "n": "teal",
                "o": "maroon",
                "p": "lime",
                "q": "silver",
                "r": "aqua",
                "s": "fuchsia",
                "t": "indigo",
                "u": "violet",
                "v": "khaki",
                "w": "turquoise",
                "x": "coral",
                "y": "gold",
                "z": "orchid",
                "0": "red",
                "1": "blue",
                "2": "green",
                "3": "orange",
                "4": "purple",
                "5": "brown",
                "6": "grey",
                "7": "pink",
                "8": "cyan",
                "9": "magenta",
            }
            self.config = {
                'color_mapping': default_color_mapping
            }
               
            
        text_frame = tk.Frame(master)
        text_frame.pack(side='left', fill='both', expand=True)
        text_scroll = tk.Scrollbar(text_frame)
        text_scroll.pack(side='right', fill='y')
        self.text = tk.Text(text_frame, yscrollcommand=text_scroll.set)
        self.text.pack(side='left', fill='both', expand=True)
        text_scroll.config(command=self.text.yview)
        self.text.bind("<Key>", self.colorize_after_event)
        self.text.bind("<Control-v>", self.colorize_after_event)
        self.setup_menu()
        self.last_line=0

    def setup_menu(self):
        menu_bar = tk.Menu(self.master)
        file_menu = tk.Menu(menu_bar, tearoff=0)
        file_menu.add_command(label="Open", command=self.open_file)
        file_menu.add_command(label="Save", command=self.save_file)
        file_menu.add_command(label="Save as Word file (docx)", command=self.save_word_file)

        menu_bar.add_cascade(label="File", menu=file_menu)
        format_menu = tk.Menu(menu_bar, tearoff=0)
        format_menu.add_checkbutton(label="Bold", command=self.toggle_bold)
        format_menu.add_checkbutton(label="Italic", command=self.toggle_italic)
        format_menu.add_checkbutton(label="Underline", command=self.toggle_underline)
        format_menu.add_command(label="Font", command=self.choose_font)
        menu_bar.add_cascade(label="Format", menu=format_menu)       
        config_menu = tk.Menu(menu_bar, tearoff=0)
        config_menu.add_command(label="Edit color mapping", command=self.edit_color_mapping)
        config_menu.add_command(label="Import color mapping", command=self.import_color_mapping)
        config_menu.add_command(label="Export color mapping", command=self.export_color_mapping)
        config_menu.add_command(label="Colorize whole text", command=self.colorize_all)
        menu_bar.add_cascade(label="Color mapping", menu=config_menu)
        self.master.config(menu=menu_bar)
    
    def choose_font(self):
        font_tuple = tkfont.Font(font=self.text['font']).actual()
        font_name = font_tuple['family']
        font_size = font_tuple['size']
        new_font = tkfont.families()
        new_font_window = tk.Toplevel(self.master)
        new_font_window.title("Choose Font")
        new_font_label = tk.Label(new_font_window, text="Choose a font:")
        new_font_label.grid(row=0, column=0, padx=5, pady=5)
        new_font_var = tk.StringVar(value=font_name)
        new_font_listbox = tk.Listbox(new_font_window, listvariable=new_font_var, height=8)
        new_font_listbox.grid(row=1, column=0, padx=5, pady=5)
        new_font_scrollbar = tk.Scrollbar(new_font_window)
        new_font_scrollbar.grid(row=1, column=1, sticky='NS')
        new_font_listbox.config(yscrollcommand=new_font_scrollbar.set)
        new_font_scrollbar.config(command=new_font_listbox.yview)
        for font in new_font:
            new_font_listbox.insert(tk.END, font)
        new_font_size_label = tk.Label(new_font_window, text="Choose a font size:")
        new_font_size_label.grid(row=2, column=0, padx=5, pady=5)
        new_font_size_var = tk.IntVar(value=font_size)
        new_font_size_spinbox = tk.Spinbox(new_font_window, from_=1, to=72, textvariable=new_font_size_var, width=5)
        new_font_size_spinbox.grid(row=3, column=0, padx=5, pady=5)
        new_font_ok_button = tk.Button(
            new_font_window, text="OK", 
            command=lambda: self.change_font(new_font_listbox.get(tk.ACTIVE), new_font_size_var.get()))
        new_font_ok_button.grid(row=4, column=0, padx=5, pady=5)

    def change_font(self, font_name, font_size):
        self.text.config(font=(font_name, font_size))
    
    def open_file(self):
        filename = filedialog.askopenfilename(filetypes=[('Text files', '*.txt')])
        if filename:
            with open(filename, "r") as file:
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, file.read())

    def save_file(self):
        filename = filedialog.asksaveasfilename(
            defaultextension='.txt',
            filetypes=[('Text files', '*.txt')])
        if filename:
            with open(filename, "w") as file:
                file.write(self.text.get(1.0, tk.END))
                
    def save_word_file(self):
        filename = filedialog.asksaveasfilename(
            defaultextension='.docx',
            filetypes=[('Text files', '*.docx')])
        if filename:
            document = Document()
            text = self.text.get(1.0, tk.END)
            paragraph = document.add_paragraph()
            for char in list(text):
                run=paragraph.add_run(char)
                font=run.font
                color = self.config['color_mapping'].get(char.lower(), "black")
                rgb = tuple((c//256 for c in self.master.winfo_rgb(color)))
                font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
            document.save(filename)
                
    def toggle_bold(self):
        bold_tag = "bold"
        if self.text.tag_ranges(bold_tag):
            self.text.tag_remove(bold_tag, "sel.first", "sel.last")
        else:
            self.text.tag_add(bold_tag, "sel.first", "sel.last")
            self.text.tag_config(bold_tag, font=("TkDefaultFont", 12, "bold"))
    
    def toggle_italic(self):
        italic_tag = "italic"
        if self.text.tag_ranges(italic_tag):
            self.text.tag_remove(italic_tag, "sel.first", "sel.last")
        else:
            self.text.tag_add(italic_tag, "sel.first", "sel.last")
            self.text.tag_config(italic_tag, font=("TkDefaultFont", 12, "italic"))
    
    def toggle_underline(self):
        underline_tag = "underline"
        if self.text.tag_ranges(underline_tag):
            self.text.tag_remove(underline_tag, "sel.first", "sel.last")
        else:
            self.text.tag_add(underline_tag, "sel.first", "sel.last")
            self.text.tag_config(underline_tag, underline=1)
      
    def save_config(self):
        print('Save config at', self.config_file)
        os.makedirs(os.path.dirname(self.config_file), exist_ok=True)
        with open(self.config_file, 'w') as f:
            json.dump(self.config, f)
    
    def edit_color_mapping(self):
        EditColorMapping(
            self.master, self.config['color_mapping'], 
            self.colorize_text, self.update_color_mapping)
        
    def update_color_mapping(self, new_color_mapping):
        print('Update color mapping.')
        self.config['color_mapping'] = new_color_mapping
        self.colorize_text()
        self.save_config()
                
    def export_color_mapping(self):
        filename = filedialog.asksaveasfilename(
            defaultextension='.json',
            initialdir=program_path,
            filetypes=[('JSON Files', '*.json')]
        )
        if filename:
            with open(filename, 'w') as f:
                json.dump(self.config['color_mapping'], f)
    
    def import_color_mapping(self):
        filename = filedialog.askopenfilename(
            initialdir=program_path,
            filetypes=[('JSON Files', '*.json')]
        )
        if filename:
            with open(filename, 'r') as f:
                self.config['color_mapping'] = json.load(f)
            self.colorize_text()
            self.save_config()
            
    def colorize_after_event(self, event):   
        color = self.config['color_mapping'].get(event.char.lower(), "black")
        self.text.tag_delete('at_cursor')
        self.text.tag_config('at_cursor', foreground=color)
        # Necessary if inserted within text:
        self.text.tag_add("at_cursor", "insert-1c", "insert")
        # Required if character is added at the end:
        self.text.tag_add("at_cursor", "insert", "insert+1c")
        
        index = self.text.index('insert')
        line_num = int(index.split('.')[0])
        self.colorize_line(line_num)
        self.colorize_line(self.last_line)
        self.last_line=line_num

    def colorize_all(self):
        self.colorize_text()
        
    def colorize_text(self, color_mapping=None):
        if not color_mapping:
            color_mapping=self.config['color_mapping'];
        for line in range(1, int(self.text.index("end").split(".")[0]) + 1):
            self.colorize_line(line,color_mapping)
            
    def colorize_line(self,line,color_mapping=None): 
        if not color_mapping:
            color_mapping=self.config['color_mapping'];
        line_start = f"{line}.0"
        line_end = f"{line}.end"
        line_text = self.text.get(line_start, line_end)
        for j, char in enumerate(line_text):
            if char:
                color = color_mapping.get(char.lower(), "black")
                tag_name = f"{line}.{j}"
                start_pos = f"{line}.{j}"
                end_pos = f"{line}.{j+1}"
                self.text.tag_delete(tag_name)
                self.text.tag_add(tag_name, start_pos, end_pos)
                self.text.tag_config(tag_name, foreground=color)
              
def main():
    root = tk.Tk()
    text_editor = TextEditor(root)
    try:
        root.mainloop()
    except KeyboardInterrupt:
        # Close the window and terminate the program
        root.destroy()

if __name__ == '__main__':
    main()